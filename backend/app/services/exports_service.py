from __future__ import annotations
import io, time
from pathlib import Path
from typing import Dict, Any, List
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from reportlab.lib.units import cm
from docx import Document
from docx.shared import RGBColor
from backend.app.core.config import settings

class ExportsService:
    def __init__(self):
        self.base = Path(settings.EXPORTS_DIR)
        (self.base / "pdf_reports").mkdir(parents=True, exist_ok=True)
        (self.base / "csv").mkdir(parents=True, exist_ok=True)
        (self.base / "parquet").mkdir(parents=True, exist_ok=True)
        (self.base / "docx_redlines").mkdir(parents=True, exist_ok=True)

    def exec_brief(self, risk: Dict[str, Any]) -> Dict[str, str]:
        # risk = {"lens":{"legal":..},"composite":..,"top_issues":[...], "renewals":[{"contract":"X","date":"2026-01-01"}]}
        lens = risk.get("lens", {})
        arr = np.array([[lens.get("legal",0), lens.get("operational",0), lens.get("regulatory",0)],
                        [lens.get("counterparty",0), lens.get("financial",0), risk.get("composite",0)]])
        fig, ax = plt.subplots(figsize=(3,2))
        im = ax.imshow(arr, vmin=0, vmax=1)
        ax.set_xticks([0,1,2]); ax.set_xticklabels(["Legal","Oper","Reg"])
        ax.set_yticks([0,1]); ax.set_yticklabels(["CP/Fin","Composite"])
        plt.colorbar(im)
        buf = io.BytesIO()
        plt.tight_layout()
        plt.savefig(buf, format="png", dpi=200)
        plt.close(fig)
        buf.seek(0)

        pdf_path = self.base / "pdf_reports" / f"executive_brief_{int(time.time())}.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=A4)
        width, height = A4
        c.setFont("Helvetica-Bold", 16)
        c.drawString(2*cm, height-2*cm, "Executive Brief — Contract Risk")
        c.setFont("Helvetica", 10)
        c.drawString(2*cm, height-3*cm, f"Composite Risk: {risk.get('composite',0):.2f}")
        img = ImageReader(buf)
        c.drawImage(img, 2*cm, height-10*cm, width=10*cm, preserveAspectRatio=True, mask='auto')
        # top 5 issues
        c.setFont("Helvetica-Bold", 12)
        c.drawString(2*cm, height-11*cm, "Top 5 Issues")
        c.setFont("Helvetica", 10)
        y = height-12*cm
        for i, issue in enumerate(risk.get("top_issues", [])[:5], start=1):
            c.drawString(2.2*cm, y, f"{i}. {issue}")
            y -= 0.8*cm
        # renewal timeline (simple list)
        c.setFont("Helvetica-Bold", 12)
        c.drawString(2*cm, y-0.5*cm, "Upcoming Renewals")
        y -= 1.2*cm
        for r in risk.get("renewals", [])[:5]:
            c.drawString(2.2*cm, y, f"- {r.get('contract')} → {r.get('date')}")
            y -= 0.7*cm
        c.showPage()
        c.save()
        return {"pdf_path": str(pdf_path)}

    def star_schema(self, rows: List[Dict[str, Any]]) -> Dict[str, str]:
        # rows: [{contract_id, counterparty, bu, lens:{...}, composite}]
        flat = []
        for r in rows:
            lens = r.get("lens", {})
            flat.append({
                "ContractID": r.get("contract_id","C-000"),
                "Counterparty": r.get("counterparty","Unknown"),
                "BU": r.get("bu","default"),
                "Legal": lens.get("legal",0.0),
                "Operational": lens.get("operational",0.0),
                "Regulatory": lens.get("regulatory",0.0),
                "CounterpartyRisk": lens.get("counterparty",0.0),
                "Financial": lens.get("financial",0.0),
                "Composite": r.get("composite",0.0),
            })
        fact = pd.DataFrame(flat)
        csv_p = self.base / "csv" / f"Fact_ContractRisk_{int(time.time())}.csv"
        pq_p  = self.base / "parquet" / f"Fact_ContractRisk_{int(time.time())}.parquet"
        fact.to_csv(csv_p, index=False)
        fact.to_parquet(pq_p, index=False)
        return {"csv": str(csv_p), "parquet": str(pq_p)}

    def redline(self, original: str, revised: str, title: str) -> Dict[str, str]:
        # DOCX with styled inserts/deletes and comments
        doc = Document()
        doc.add_heading(title, level=1)
        import difflib
        s = difflib.ndiff(original.split(), revised.split())
        p = doc.add_paragraph()
        for token in s:
            code = token[:2]
            word = token[2:]
            run = p.add_run(word + " ")
            if code == "- ":
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # red delete
                run.italic = True
            elif code == "+ ":
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # green insert
                run.bold = True
        doc.add_paragraph("Note: Visual redline (styled). Track-changes not available in this demo.")
        docx_path = self.base / "docx_redlines" / f"redline_{int(time.time())}.docx"
        doc.save(str(docx_path))

        # Side-by-side diff PDF
        pdf_path = self.base / "pdf_reports" / f"diff_{int(time.time())}.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=A4)
        width, height = A4
        c.setFont("Helvetica-Bold", 12)
        c.drawString(2*cm, height-2*cm, "Original")
        c.drawString(11*cm, height-2*cm, "Suggested (not legal advice)")
        c.setFont("Helvetica", 9)
        # simple text wrap
        def wrap(txt, max_chars=70):
            words = txt.split()
            lines=[]; cur=""
            for w in words:
                if len(cur)+len(w)+1>max_chars:
                    lines.append(cur); cur=w
                else:
                    cur = (cur+" "+w).strip()
            if cur: lines.append(cur)
            return lines
        y1 = height-3*cm
        for ln in wrap(original, 70):
            c.drawString(2*cm, y1, ln); y1 -= 0.5*cm
        y2 = height-3*cm
        for ln in wrap(revised, 70):
            c.drawString(11*cm, y2, ln); y2 -= 0.5*cm
        c.showPage(); c.save()

        return {"docx": str(docx_path), "pdf": str(pdf_path)}
