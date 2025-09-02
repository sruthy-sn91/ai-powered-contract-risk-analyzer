from __future__ import annotations
import io
import re
import math
import logging
from pathlib import Path
from typing import List, Optional, Tuple, Dict
from fastapi import UploadFile
from PIL import Image
import fitz  # PyMuPDF
import pdfplumber
import pytesseract
from langdetect import detect as lang_detect
from backend.app.schemas.extraction import (
    ParsedPage,
    ClauseSegment,
    ParsingResult,
    DetectedMeta,
)
from backend.app.core.path_resolver import index_dir

WATERMARK_TERMS = ["CONFIDENTIAL", "DRAFT", "SAMPLE", "VOID", "WATERMARK"]
GOV_LAW_PAT = re.compile(r"(?i)\bgoverned by(?: the laws of)? ([A-Za-z ,&.-]+)")
JURIS_PAT = re.compile(r"(?i)\bexclusive jurisdiction of ([A-Za-z ,&.-]+)|\bcourts of ([A-Za-z ,&.-]+)")
CURRENCY_SYMBOLS = {
    "$": "USD", "€": "EUR", "£": "GBP", "¥": "JPY", "₹": "INR", "C$": "CAD", "A$": "AUD",
}
CURRENCY_CODES = {"USD","EUR","GBP","JPY","INR","CAD","AUD","CHF","CNY","HKD","SGD","SEK","NOK","DKK"}

def _bytes_to_path(tmp_dir: Path, filename: str, data: bytes) -> Path:
    tmp_dir.mkdir(parents=True, exist_ok=True)
    p = tmp_dir / filename
    p.write_bytes(data)
    return p

def _extract_with_fitz(pdf_path: Path) -> Tuple[List[ParsedPage], List[List[str]]]:
    pages: List[ParsedPage] = []
    tables_flags: List[List[str]] = []

    with fitz.open(str(pdf_path)) as doc:
        for i, page in enumerate(doc, start=1):
            rot = int(page.rotation or 0)
            blocks = page.get_text("blocks")
            lefts = [b[0] for b in blocks if isinstance(b[4], str) and b[4].strip()]
            if lefts:
                median_x = sorted(lefts)[len(lefts)//2]
            else:
                median_x = page.rect.width / 2

            left_col, right_col = [], []
            for b in blocks:
                txt = b[4] if len(b) > 4 and isinstance(b[4], str) else ""
                if not txt.strip():
                    continue
                (left_col if b[0] < median_x else right_col).append((b[1], b[0], txt))

            left_col.sort(key=lambda x: (x[0], x[1]))
            right_col.sort(key=lambda x: (x[0], x[1]))
            ordered_text = "\n".join([t[2] for t in left_col + right_col]).strip()

            wm = []
            upper = ordered_text.upper()
            for term in WATERMARK_TERMS:
                if term in upper:
                    wm.append(term)

            area = page.rect.width * page.rect.height or 1.0
            density = min(1.0, max(0.0, len(ordered_text) / max(500.0, area / 5000.0)))

            pages.append(ParsedPage(
                page_number=i,
                text=ordered_text,
                ocr_used=False,
                rotation=rot,
                has_tables=False,
                watermarks=wm,
                quality_score=float(f"{density:.3f}")
            ))
    try:
        with pdfplumber.open(str(pdf_path)) as pl:
            for idx, p in enumerate(pl.pages, start=1):
                page_tables = p.extract_tables() or []
                tables_flags.append([str(len(page_tables))])
                if 1 <= idx <= len(pages):
                    pages[idx-1].has_tables = bool(page_tables)
    except Exception as e:
        logging.getLogger(__name__).warning(f"pdfplumber failed: {e}")

    return pages, tables_flags

def _ocr_low_text_pages(pdf_path: Path, pages: List[ParsedPage]) -> None:
    logger = logging.getLogger(__name__)
    try:
        with fitz.open(str(pdf_path)) as doc:
            for idx, page in enumerate(doc, start=1):
                pp = pages[idx-1]
                if pp.quality_score >= 0.4 and len(pp.text) > 100:
                    continue
                pix = page.get_pixmap(dpi=300, annots=False)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                try:
                    txt = pytesseract.image_to_string(img)
                except Exception as ocr_e:
                    logger.warning(f"OCR failed on page {idx}: {ocr_e}")
                    continue
                if txt and len(txt) > len(pp.text) * 1.2:
                    pp.text = txt
                    pp.ocr_used = True
                    pp.quality_score = min(1.0, max(0.0, len(txt) / 1500.0))
    except Exception as e:
        logger.warning(f"OCR routine error: {e}")

def _docx_to_text(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    paras = []
    for p in doc.paragraphs:
        paras.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            paras.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(s for s in paras if s is not None)

HEADING_PAT = re.compile(
    r"(?m)^(?:\s*((?:\d{1,2}\.){1,4}\d{0,2}|[A-Z][A-Z0-9 \-]{3,}))\s*$"
)

def _segment_clauses(text: str):
    matches = list(HEADING_PAT.finditer(text))
    from backend.app.schemas.extraction import ClauseSegment  # local import
    segments = []
    if not matches:
        pos = 0
        parts = [p for p in text.split("\n\n") if p.strip()]
        for i, p in enumerate(parts, start=1):
            start = text.find(p, pos)
            end = start + len(p)
            segments.append(ClauseSegment(id=f"C{i}", title=None, heading=None, start=start, end=end))
            pos = end
        return segments

    for i, m in enumerate(matches, start=1):
        start = m.start()
        end = matches[i].start() if i < len(matches) else len(text)
        heading = m.group(1).strip() if m.group(1) else None
        title = heading
        segments.append(ClauseSegment(id=f"C{i}", title=title, heading=heading, start=start, end=end))
    return segments

def _detect_meta(text: str):
    from backend.app.schemas.extraction import DetectedMeta
    meta = DetectedMeta()
    for m in re.compile(r"(?i)\bgoverned by(?: the laws of)? ([A-Za-z ,&.-]+)").finditer(text):
        g = m.group(1)
        if g:
            meta.governing_law.append(g.strip(",. ").strip())
    for m in re.compile(r"(?i)\bexclusive jurisdiction of ([A-Za-z ,&.-]+)|\bcourts of ([A-Za-z ,&.-]+)").finditer(text):
        cand = m.group(1) or m.group(2)
        if cand:
            meta.jurisdiction.append(cand.strip(",. ").strip())
    try:
        from langdetect import detect as lang_detect
        lang = lang_detect(text[:2000] if len(text) > 2000 else text)
        meta.languages.append(lang)
    except Exception:
        pass
    upp = text.upper()
    for sym, code in {"$":"USD","€":"EUR","£":"GBP","¥":"JPY","₹":"INR","C$":"CAD","A$":"AUD"}.items():
        if sym in text and code not in meta.currencies:
            meta.currencies.append(code)
    for code in {"USD","EUR","GBP","JPY","INR","CAD","AUD","CHF","CNY","HKD","SGD","SEK","NOK","DKK"}:
        if code in upp and code not in meta.currencies:
            meta.currencies.append(code)
    return meta

class ParsingOCRService:
    async def parse(self, file: Optional[UploadFile], path: Optional[str]):
        tmp = index_dir() / "_tmp"
        if file:
            data = await file.read()
            src_path = _bytes_to_path(tmp, file.filename or "upload.bin", data)
        else:
            src_path = Path(path).expanduser().resolve()

        text_all = ""
        pages = []

        if src_path.suffix.lower() == ".pdf":
            pages, _ = _extract_with_fitz(src_path)
            _ocr_low_text_pages(src_path, pages)
            for p in pages:
                text_all += (p.text or "") + "\n\n"
        elif src_path.suffix.lower() in {".docx"}:
            text_all = _docx_to_text(src_path)
            pages = [ParsedPage(page_number=1, text=text_all, ocr_used=False, rotation=0,
                                has_tables=False, watermarks=[], quality_score=1.0)]
        else:
            text_all = src_path.read_text(encoding="utf-8", errors="ignore")
            pages = [ParsedPage(page_number=1, text=text_all, ocr_used=False, rotation=0,
                                has_tables=False, watermarks=[], quality_score=1.0)]

        normalized = re.sub(r"[ \t]+", " ", text_all).replace("\r", "")
        clauses = _segment_clauses(normalized)
        meta = _detect_meta(normalized)

        return ParsingResult(
            pages=pages,
            normalized_text=normalized,
            clauses=clauses,
            meta=meta
        )
